#!/usr/bin/env python3
"""
summarize.py — 用 Gemini API 生成中文文章纪要
================================================
读取 fetch_fulltext.py 生成的 CSV，
对每篇文章调用 Gemini 生成结构化中文摘要，
输出新的 CSV 文件。

用法:
    python summarize.py --input output/gcc_fulltext_20260323_1506.csv --key YOUR_API_KEY
    python summarize.py --input output/gcc_fulltext_20260323_1506.csv --key YOUR_API_KEY --test
"""

import requests
import csv
import time
import argparse
import logging
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("summarize.log", encoding="utf-8"),
    ],
)
log = logging.getLogger("summarize")

# Gemini API 配置
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"
DELAY = 4.0  # 每次请求间隔，避免超过每分钟15次限制


# ══════════════════════════════════════════════
# Prompt 模板
# ══════════════════════════════════════════════
PROMPT_TEMPLATE = """你是一位为中国政府部门提供参考的国际问题研究员，专注于海湾合作委员会（GCC）地区的政治、经济与金融动态。

请根据以下英文文章，撰写一份供政府官员参阅的中文研究纪要。语言要求：简洁、正式、结论明确，避免学术腔，直接说明关键信息。

文章信息：
- 来源机构：{source}
- 文章标题：{title}
- 发布日期：{date}

文章正文：
{full_text}

请按以下结构输出（总字数400-500字，每个板块简明扼要）：

【核心议题】
一句话说明本文研究什么问题。

【主要判断】
列出2-4条文章的核心观点或结论，每条不超过50字，语气肯定。

【对GCC地区的影响】
说明上述判断对海湾六国（沙特、阿联酋、卡塔尔、巴林、科威特、阿曼）政治、经济或安全格局的实际影响（100字以内）。

【对华关联】
说明本文议题与中国利益的关联，例如：能源供应、投资合作、一带一路、地区稳定等方面的影响（如文章未涉及对华内容，可根据议题性质作简要推断，50字以内）。

【关键数据或事件】
列出文章中1-3条重要数据、协议或具体事件（如无则省略）。

【来源背景】
一句话说明发布机构的性质和立场（如：美国智库、卡塔尔官方背景研究机构等）。
"""


# ══════════════════════════════════════════════
# 调用 Gemini API
# ══════════════════════════════════════════════
def call_gemini(prompt: str, api_key: str) -> str:
    url = f"{GEMINI_URL}?key={api_key}"
    payload = {
        "contents": [
            {
                "parts": [{"text": prompt}]
            }
        ],
        "generationConfig": {
            "temperature": 0.3,      # 低温度，输出更稳定
            "maxOutputTokens": 4096,
        }
    }
    for attempt in range(3):
        try:
            r = requests.post(url, json=payload, timeout=30)
            if r.status_code == 429:
                wait = 30 * (attempt + 1)
                log.warning(f"  请求频率超限，等待{wait}秒后重试...")
                time.sleep(wait)
                continue
            r.raise_for_status()
            data = r.json()
            return data["candidates"][0]["content"]["parts"][0]["text"]
        except requests.RequestException as e:
            log.error(f"  API请求失败: {type(e).__name__}")
            time.sleep(10)
        except (KeyError, IndexError) as e:
            log.error(f"  API返回格式异常: {e}")
            return ""
    return ""


# ══════════════════════════════════════════════
# 主程序
# ══════════════════════════════════════════════
def add_paragraph(doc, text, bold=False, size=11, color=None, space_before=0, space_after=6, align=None):
    """添加段落的辅助函数"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def generate_word(results: list, output_path: Path, ts: str):
    """用 python-docx 生成格式化的Word文档"""
    doc = Document()

    # 页面设置：A4
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 封面标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("GCC智库研究动态集锦")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # 日期
    date_str = f"{ts[:4]}年{ts[4:6]}月{ts[6:8]}日"
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(18)
    date_run = date_p.add_run(f"编制日期：{date_str}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 每篇文章
    valid = [(i, r) for i, r in enumerate(results) if r.get("summary_zh", "").strip()]
    for idx, (i, item) in enumerate(valid):

        # 文章标题
        t = f"{idx+1}、{item['title']}"
        h = doc.add_heading(t, level=1)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(4)
        for run in h.runs:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        # 来源 / 日期
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(8)
        r1 = meta_p.add_run("来源：")
        r1.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r2 = meta_p.add_run(item.get("source_name", "") + "    ")
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r3 = meta_p.add_run("日期：")
        r3.bold = True
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r4 = meta_p.add_run(item.get("date", "未知"))
        r4.font.size = Pt(9)
        r4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 摘要正文
        summary = item.get("summary_zh", "")
        for line in summary.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue
            is_header = stripped.startswith("【") and "】" in stripped
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6 if is_header else 0)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped)
            run.bold = is_header
            run.font.size = Pt(10.5 if is_header else 10)
            if is_header:
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        # 原文链接
        url_p = doc.add_paragraph()
        url_p.paragraph_format.space_before = Pt(4)
        url_p.paragraph_format.space_after = Pt(4)
        ur1 = url_p.add_run("原文链接：")
        ur1.bold = True
        ur1.font.size = Pt(9)
        ur1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        ur2 = url_p.add_run(item.get("url", ""))
        ur2.font.size = Pt(9)
        ur2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        # 分割线（用段落下边框）
        if idx < len(valid) - 1:
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(12)
            sep.paragraph_format.space_after = Pt(0)
            pPr = sep._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)

    try:
        doc.save(str(output_path))
        log.info(f"Word文档生成成功: {output_path}")
    except Exception as e:
        log.error(f"Word保存失败: {e}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="输入CSV（gcc_fulltext_*.csv）")
    parser.add_argument("--key", required=True, help="Gemini API Key")
    parser.add_argument("--test", action="store_true", help="测试模式，只处理前3篇")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"文件不存在: {input_path}")
        return

    # 读取输入
    with open(input_path, encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    log.info(f"读取 {len(rows)} 篇文章")

    if args.test:
        rows = rows[:3]
        log.info("测试模式：只处理前3篇")

    # 过滤掉没有正文的文章
    valid_rows = [r for r in rows if int(r.get("word_count", 0)) > 50]
    skipped = len(rows) - len(valid_rows)
    if skipped:
        log.info(f"跳过 {skipped} 篇（无正文）")

    # 输出文件
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = input_path.parent / f"gcc_summaries_{ts}.csv"
    fields = ["source_name", "title", "date", "url",
              "gcc_keywords", "word_count", "summary_zh", "summarized_at"]

    results = []
    for i, row in enumerate(valid_rows, 1):
        title = row.get("title", "")
        source = row.get("source_name", "")
        date = row.get("date", "")
        full_text = row.get("full_text", "")

        log.info(f"[{i}/{len(valid_rows)}] {title[:60]}...")

        # 如果正文太长，截取前4000词（避免超出token限制）
        words = full_text.split()
        if len(words) > 4000:
            full_text = " ".join(words[:4000]) + "..."
            log.info(f"  正文过长，截取前4000词")

        prompt = PROMPT_TEMPLATE.format(
            source=source,
            title=title,
            date=date,
            full_text=full_text,
        )

        summary = call_gemini(prompt, args.key)

        if summary:
            log.info(f"  ✅ 摘要生成成功（{len(summary)}字符）")
        else:
            log.warning(f"  ❌ 摘要生成失败")

        results.append({
            "source_name": source,
            "title": title,
            "date": date,
            "url": row.get("url", ""),
            "gcc_keywords": row.get("gcc_keywords", ""),
            "word_count": row.get("word_count", ""),
            "summary_zh": summary,
            "summarized_at": datetime.now().isoformat(),
        })

        time.sleep(DELAY)

    # 先保存CSV备份
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fields, quoting=csv.QUOTE_ALL)
        writer.writeheader()
        writer.writerows(results)

    # 生成Word文档
    docx_path = output_path.with_suffix(".docx")
    generate_word(results, docx_path, ts)

    success = sum(1 for r in results if r["summary_zh"])
    log.info(f"\n完成：{success}/{len(results)} 篇成功生成摘要")
    log.info(f"CSV备份：{output_path}")
    log.info(f"Word文档：{docx_path}")
    print(f"\n✅ 完成！Word文档：{docx_path}")


if __name__ == "__main__":
    main()
